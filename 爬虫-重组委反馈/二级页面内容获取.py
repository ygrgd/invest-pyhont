# -*- coding: utf-8 -*-
"""
Created on Sat Jan 13 13:11:29 2018

@author: LZ
"""


from urllib.request import urlopen 
from bs4 import BeautifulSoup  as bs
import os  
import re  
import sys 
from lxml import etree
from docx import Document as doc
from docx.shared import Inches as inc





#定义网址
url = "http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html"
#打开网址传输到file
file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html')
#从file生成b0
b0 = bs(file,'lxml')
#找出b0中所有的div class =c u
a = b0.find_all('div',{'class':"Custom_UnionStyle"})
#打印其中的文字并过滤掉转换符
print(a[0].text.replace('\xa0'*8,'\n\n'))
# 生成反馈意见 fkjy
fkyj = a[0].text.replace('\xa0'*8,'\n\n')
print(fkyj)


test=(fkyj.find('契约型私募基金'),fkyj.find('资产管理计划'),fkyj.find('信托计划')\
,fkyj.find('有限合伙'),fkyj.find('契约基金'),fkyj.find('资管计划'))
a = (-1, -1, -1, -1, -1, -1)





if '股东' in fkyj:
    print('Y')
else:
    print('N')
    
    
    


#创建doc文件
document = doc()
#添加一级标题为公司名称
document.add_heading('福建三钢闽光股份有限公司', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#新开一页
document.add_page_break()
#添加下一个一级标题
document.add_heading('福建三钢闽光股份有限公司2', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#保存反馈意见
document.save("反馈意见.docx")