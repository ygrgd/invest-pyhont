# -*- coding: utf-8 -*-
"""
Created on Sat Jan 13 15:30:16 2018

@author: LZ
"""

from urllib.request import urlopen 
from bs4 import BeautifulSoup  as bs
import os  
import re  
import sys 
from lxml import etree
from docx import Document as doc
from docx.shared import Inches as inc
import pandas as pd
import csv
import numpy as np  

# =============================================================================
# 从文件中读取
# =============================================================================
csvfile = open('name_add.csv', 'r',encoding='utf-8')
data = []
for line in csvfile:
    data.append(list(line.strip().split('\t')))

document = doc()
    
for i in range(1,len(data)):
    xx = str(data[i])
    xxx = xx.split(',')
    name = xxx[0][2:]
    add = xxx[1][:-2]
    print(add)
   # file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html')
    #file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/./201801/t20180112_332401.html')
    file = urlopen(add)
    print('现在正在下载第'+str(i+1)+'个')
    print(add)
    b0 = bs(file,'lxml')
    a = b0.find_all('div',{'class':"Custom_UnionStyle"})
    fkyj = a[0].text.replace('\xa0'*8,'\n\n')
    document.add_heading(name, level=1)
    document.add_paragraph(fkyj)
    document.add_page_break()
    print(name)
    
document.save("反馈意见.docx")
#添加一级标题为公司名称

#添加段落为反馈意见

for i in range(163,len(data)):
    xx = str(data[i])
    xxx = xx.split(',')
    name = xxx[0][2:]
    add = xxx[1][:-2]
    print(add)
   # file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html')
    #file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/./201801/t20180112_332401.html')
    file = urlopen(add)
    print('现在正在下载第'+str(i+1)+'个')
    print(add)
    b0 = bs(file,'lxml')
    a = b0.find_all('div',{'class':"Custom_UnionStyle"})
    fkyj = a[0].text.replace('\xa0'*8,'\n\n')
    document.add_heading(name, level=1)
    document.add_paragraph(fkyj)
    document.add_page_break()
    print(name)
    
document.save("反馈意见.docx")

# =============================================================================
# 生成不连续的列表
# =============================================================================
a = []
for i in range(1,10) :
    a.append(i)

for i in range(12,20) :
    a.append(i)











#新开一页
document.add_page_break()
#添加下一个一级标题
document.add_heading('福建三钢闽光股份有限公司2', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#保存反馈意见
document.save("反馈意见.docx")
    
    
    
    
    
    
    

xx = str(data[1])
xxx=xx.split(',')
xxx[0]
xxx[1]    
    
    
    
    

'''
data[1].strip().split(',')

 
import csv  
with open('name_add.csv','r',encoding='utf-8') as myFile:  
    lines=csv.reader(myFile)  
    for line in lines:
        for x in line:
            xx =x.strip().split(',')
            data.append(xx)
            print(xx)
        print (x)  
'''

#定义网址
url = "http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html"
#打开网址传输到file
file = urlopen('add')
#从file生成b0
b0 = bs(file,'lxml')
#找出b0中所有的div class =c u
a = b0.find_all('div',{'class':"Custom_UnionStyle"})
#打印其中的文字并过滤掉转换符
print(a[0].text.replace('\xa0'*8,'\n\n'))
# 生成反馈意见 fkjy
fkyj = a[0].text.replace('\xa0'*8,'\n\n')

print(fkyj)

#创建doc文件
document = doc()
#添加一级标题为公司名称
document.add_heading('福建三钢闽光股份有限公司', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#新开一页
document.add_page_break()
#添加下一个一级标题
document.add_heading('福建三钢闽光股份有限公司2', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#保存反馈意见
document.save("反馈意见.docx")