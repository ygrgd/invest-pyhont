# -*- coding: utf-8 -*-
"""
Created on Sat Jan 13 15:21:04 2018

@author: LZ
"""

from urllib.request import urlopen 
from bs4 import BeautifulSoup  as bs
import os  
import re  
import sys 
from lxml import etree
from urllib.request import urlopen 
from bs4 import BeautifulSoup  as bs
import os  
import re  
import sys 
from lxml import etree
from docx import Document as doc
from docx.shared import Inches as inc
import pandas as pd
import csv
'''
http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/index.htm

http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/index_1.htm

一共到第13页 为17年度数据

'''
#生成列表 urladd 里面有第1到12个网址
url0 = 'http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/index.htm'
urladd = [url0]
urlstart = 'http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/index_'
for i in range(1,25):
    urladd.append(urlstart+str(i)+'.htm')
    print (urladd)

#定义网页的开头
server = 'http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/'
#生成bs对象



articles = []
for url in urladd:
    b0 = bs(urlopen(url),'lxml')
    a_df = b0.find(id = "myul")
    a = a_df.find_all('a')
    for each in a:
        name = each.string
        add = server + each.get('href')
        date = each.find('span').get_text()
        articles.append([name,date,add])
        print (name ,date,add)
#w 写入 a添加 newline="" 可以有效删除空行      
with open('name_add.csv', 'w',newline='',encoding ="utf-8") as f:
    writer = csv.writer(f)
    writer.writerow(['公司名字','网址'])
    for row in articles:
        writer.writerow(row)


###############################################################



#定义网址
url = "http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html"
#打开网址传输到file
file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html')
#从file生成b0
b0 = bs(file,'lxml')
#找出b0中所有的div class =c u
a = b0.find_all('div',{'class':"Custom_UnionStyle"})
#打印其中的文字并过滤掉转换符
print(a[0].text.replace('\xa0'*8,'\n\n'))
# 生成反馈意见 fkjy
fkyj = a[0].text.replace('\xa0'*8,'\n\n')

print(fkyj)

#创建doc文件
document = doc()
#添加一级标题为公司名称
document.add_heading('福建三钢闽光股份有限公司', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#新开一页
document.add_page_break()
#添加下一个一级标题
document.add_heading('福建三钢闽光股份有限公司2', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#保存反馈意见
document.save("反馈意见.docx")




















