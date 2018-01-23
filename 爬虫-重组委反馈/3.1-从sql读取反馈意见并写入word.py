# -*- coding: utf-8 -*-
"""
Created on Tue Jan 23 23:47:12 2018

@author: LZ
"""


import requests
from bs4 import BeautifulSoup  as bs
from docx import Document as doc
import pandas as pd
import sqlalchemy as say
from sqlalchemy import create_engine
import MySQLdb
import time
time.localtime()
now = time.strftime('%Y-%m-%d_%H-%M',time.localtime(time.time()))

engine = create_engine('mysql://root:@127.0.0.1/csrc?charset=utf8',echo = False)  # 定义引擎
engine.execute('use csrc')


conn = engine.connect()

content = engine.execute('select content from bgczfkyj').fetchone()
url = engine.execute('select url from bgczfkyj').fetchall()
title = engine.execute('select title from bgczfkyj').fetchall()
# =============================================================================
# 下面这个方法可行，但是没有换行。估计还是得用 orm方式获取内容才行
# =============================================================================
document = doc()

for a in range(0,496):
    i = 496-a
    document.add_heading(title[i][0],level=1)
    document.add_paragraph(content[i])
    document.add_page_break()
    print(title[i][0])
    

file_name = '反馈意见'+now+'.docx'

document.save(file_name)


# =============================================================================
# 尝试用pd 可以用
# =============================================================================

df = pd.read_sql('bgczfkyj',engine)
df = df.sort_index(ascending=False)
df = df.set_index('index')

df.title[0]
df.content[0]
df.set_index('index')


document = doc()

for i in range(0,496):
    document.add_heading(df.title[i],level=1)
    document.add_paragraph(df.content[i])
#    document.add_page_break()
    print(df.title[i])




now = time.strftime('%Y-%m-%d_%H-%M',time.localtime(time.time()))
file_name = '反馈意见'+now+'.docx'

document.save(file_name)




#创建doc文件
document = doc()
#添加一级标题为公司名称
document.add_heading('福建三钢闽光股份有限公司', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#新开一页
document.add_page_break()
#添加下一个一级标题
document.add_heading('福建三钢闽光股份有限公司2', level=1)
#添加段落为反馈意见
document.add_paragraph(fkyj)
#保存反馈意见
