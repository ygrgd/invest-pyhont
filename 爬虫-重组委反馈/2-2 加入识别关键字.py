# -*- coding: utf-8 -*-
"""
Created on Sat Jan 13 19:59:20 2018

@author: LZ
"""



from urllib.request import urlopen 
from bs4 import BeautifulSoup  as bs
import os  
import re  
import sys 
from lxml import etree
from docx import Document as doc
from docx.shared import Inches as inc
import pandas as pd
import csv
import numpy as np  


csvfile = open('E:\git\python\爬虫-重组委公告\name_add.csv', 'r',encoding='utf-8')
data = []
for line in csvfile:
    data.append(list(line.strip().split('\t')))

document = doc()
    
for i in range(1,len(data)):
    xx = str(data[i])
    xxx = xx.split(',')
    name = xxx[0][2:]
    add = xxx[1][:-2]
    print(add)
   # file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html')
    #file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/./201801/t20180112_332401.html')
    file = urlopen(add)
    print('现在正在下载第'+str(i+1)+'个')
    print(add)
    b0 = bs(file,'lxml')
    a = b0.find_all('div',{'class':"Custom_UnionStyle"})
    fkyj = a[0].text.replace('\xa0'*8,'\n\n')
    test=(fkyj.find('矿'))
    if test == (-1):
        print('没有要找的关键字')
    else:
        document.add_heading(name, level=1)
        document.add_paragraph(fkyj)
        document.add_page_break()
        print(name)
    
document.save("反馈意见-矿.docx")
#添加一级标题为公司名称

#添加段落为反馈意见

for i in range(163,len(data)):
    xx = str(data[i])
    xxx = xx.split(',')
    name = xxx[0][2:]
    add = xxx[1][:-2]
    print(add)
   # file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/201801/t20180112_332401.html')
    #file = urlopen('http://www.csrc.gov.cn/pub/newsite/ssgsjgb/bgczfkyj/./201801/t20180112_332401.html')
    file = urlopen(add)
    print('现在正在下载第'+str(i+1)+'个')
    print(add)
    b0 = bs(file,'lxml')
    a = b0.find_all('div',{'class':"Custom_UnionStyle"})
    fkyj = a[0].text.replace('\xa0'*8,'\n\n')
    test=(fkyj.find('矿'))
    if test == (-1):
        print('没有要找的关键字')
    else:
        document.add_heading(name, level=1)
        document.add_paragraph(fkyj)
        document.add_page_break()
        print(name)
    
document.save("反馈意见-矿.docx")



'''
test=(fkyj.find('契约型私募基金'),fkyj.find('资产管理计划'),fkyj.find('信托计划')\
,fkyj.find('有限合伙'),fkyj.find('契约基金'),fkyj.find('资管计划'))
a = (-1, -1, -1, -1, -1, -1)
'''








